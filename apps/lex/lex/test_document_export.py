from __future__ import annotations

import io

import frappe
from frappe.tests.utils import FrappeTestCase

from lex.document_export import ExportOptions, build_export_files, normalize_export_options, parse_markdown_blocks


SAMPLE_MARKDOWN = """# Contract Review Report

## Executive Summary

The agreement creates **material obligations** and should be approved subject to the controls below.

- Cap aggregate liability at the annual fees.
- Preserve confidentiality for five years.

| Risk | Rating | Recommendation |
|---|---|---|
| Uncapped indemnity | High | Add a negotiated cap |
| Governing law | Medium | Confirm India jurisdiction |

### Decision

Proceed after the high-risk clause is remediated.
"""


class TestDocumentExport(FrappeTestCase):
	def test_markdown_parser_keeps_structured_blocks(self):
		blocks = parse_markdown_blocks(SAMPLE_MARKDOWN)
		self.assertEqual(blocks[0], {"type": "h1", "text": "Contract Review Report"})
		self.assertTrue(any(block["type"] == "bullet" for block in blocks))
		self.assertTrue(any(block["type"] == "table" and len(block["rows"]) == 3 for block in blocks))

	def test_advanced_options_are_validated(self):
		options = normalize_export_options(
			output_format="Both",
			document_title="MSA / Risk Report",
			page_size="Letter",
			document_style="Executive Brief",
			confidentiality_label="Internal Use Only",
			include_cover_page=0,
			include_metadata=1,
			include_page_numbers=0,
		)
		self.assertEqual(options.output_format, "Both")
		self.assertEqual(options.page_size, "Letter")
		self.assertFalse(options.include_cover_page)

		with self.assertRaises(frappe.ValidationError):
			normalize_export_options(output_format="HTML")

	def test_pdf_and_docx_packages_are_generated(self):
		metadata = {
			"subtitle": "Contract Review - Final Deliverable",
			"client_name": "Lexocrates Demo Client Pvt. Ltd.",
			"matter_title": "Corporate Contract Review",
			"job_id": "LPOJ-TEST-00001",
			"job_label": "LPOJ-TEST-00001 - Master Services Agreement Review",
			"version_label": "v1.0",
			"generated_label": "24 August 2026, 15:30 IST",
			"generated_by": "Administrator",
		}
		options = ExportOptions(output_format="Both", document_title="Master Services Agreement Review")
		outputs = build_export_files(SAMPLE_MARKDOWN, metadata, options)

		self.assertEqual(set(outputs), {"PDF", "DOCX"})
		self.assertTrue(outputs["PDF"].startswith(b"%PDF-"))
		self.assertTrue(outputs["DOCX"].startswith(b"PK"))

		from pypdf import PdfReader
		from docx import Document

		pdf_reader = PdfReader(io.BytesIO(outputs["PDF"]))
		self.assertGreaterEqual(len(pdf_reader.pages), 2)
		self.assertIn("Contract Review Report", "\n".join(page.extract_text() or "" for page in pdf_reader.pages))

		docx = Document(io.BytesIO(outputs["DOCX"]))
		text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
		self.assertIn("Master Services Agreement Review", text)
		self.assertIn("Contract Review Report", text)
		self.assertGreaterEqual(len(docx.tables), 2)
