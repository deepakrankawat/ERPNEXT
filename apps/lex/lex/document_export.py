from __future__ import annotations

import html
import io
import os
import re
from dataclasses import dataclass
from datetime import datetime

import frappe
from frappe import _


EXPORT_FORMATS = {"PDF", "DOCX", "Both"}
PAGE_SIZES = {"A4", "Letter"}
DOCUMENT_STYLES = {"Legal Professional", "Executive Brief", "Plain"}
CONFIDENTIALITY_LABELS = {"Privileged & Confidential", "Confidential", "Internal Use Only", "None"}


@dataclass(frozen=True)
class ExportOptions:
	output_format: str = "PDF"
	document_title: str = "Legal Operations Deliverable"
	page_size: str = "A4"
	document_style: str = "Legal Professional"
	confidentiality_label: str = "Privileged & Confidential"
	include_cover_page: bool = True
	include_metadata: bool = True
	include_page_numbers: bool = True


def normalize_export_options(
	*,
	output_format: str | None = None,
	document_title: str | None = None,
	page_size: str | None = None,
	document_style: str | None = None,
	confidentiality_label: str | None = None,
	include_cover_page: int | bool = 1,
	include_metadata: int | bool = 1,
	include_page_numbers: int | bool = 1,
) -> ExportOptions:
	output_format = (output_format or "PDF").strip()
	page_size = (page_size or "A4").strip()
	document_style = (document_style or "Legal Professional").strip()
	confidentiality_label = (confidentiality_label or "Privileged & Confidential").strip()
	if output_format not in EXPORT_FORMATS:
		frappe.throw(_("Output format must be PDF, DOCX, or Both."), frappe.ValidationError)
	if page_size not in PAGE_SIZES:
		frappe.throw(_("Page size must be A4 or Letter."), frappe.ValidationError)
	if document_style not in DOCUMENT_STYLES:
		frappe.throw(_("Unsupported document style."), frappe.ValidationError)
	if confidentiality_label not in CONFIDENTIALITY_LABELS:
		frappe.throw(_("Unsupported confidentiality label."), frappe.ValidationError)
	clean_title = re.sub(r"[\x00-\x1f]+", " ", document_title or "Legal Operations Deliverable").strip()
	return ExportOptions(
		output_format=output_format,
		document_title=clean_title[:180] or "Legal Operations Deliverable",
		page_size=page_size,
		document_style=document_style,
		confidentiality_label=confidentiality_label,
		include_cover_page=bool(int(include_cover_page or 0)),
		include_metadata=bool(int(include_metadata or 0)),
		include_page_numbers=bool(int(include_page_numbers or 0)),
	)


def build_export_files(markdown_text: str, metadata: dict, options: ExportOptions) -> dict[str, bytes]:
	if not (markdown_text or "").strip():
		frappe.throw(_("Deliverable content is empty."), frappe.ValidationError)
	outputs = {}
	if options.output_format in {"PDF", "Both"}:
		outputs["PDF"] = generate_pdf_bytes(markdown_text, metadata, options)
	if options.output_format in {"DOCX", "Both"}:
		outputs["DOCX"] = generate_docx_bytes(markdown_text, metadata, options)
	return outputs


def generate_pdf_bytes(markdown_text: str, metadata: dict, options: ExportOptions) -> bytes:
	try:
		from reportlab.lib import colors
		from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
		from reportlab.lib.pagesizes import A4, LETTER
		from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
		from reportlab.lib.units import inch
		from reportlab.platypus import (
			Image,
			PageBreak,
			Paragraph,
			SimpleDocTemplate,
			Spacer,
			Table,
			TableStyle,
		)
	except ImportError:
		frappe.throw(_("PDF export dependency is unavailable. Install reportlab."), frappe.ValidationError)

	page = A4 if options.page_size == "A4" else LETTER
	buffer = io.BytesIO()
	doc = SimpleDocTemplate(
		buffer,
		pagesize=page,
		rightMargin=0.78 * inch,
		leftMargin=0.78 * inch,
		topMargin=0.82 * inch,
		bottomMargin=0.72 * inch,
		title=options.document_title,
		author="Lexocrates Legal Services Pvt. Ltd.",
		subject=f"LPO Job {metadata.get('job_id') or ''}",
	)
	styles = getSampleStyleSheet()
	ink = colors.HexColor("#22344F")
	muted = colors.HexColor("#607086")
	accent = colors.HexColor("#1976D2")
	body_font = "Times-Roman" if options.document_style == "Legal Professional" else "Helvetica"
	body_size = 10.5 if options.document_style != "Plain" else 11
	body_leading = 15 if options.document_style == "Legal Professional" else 14.5
	style_map = {
		"body": ParagraphStyle("LexBody", parent=styles["BodyText"], fontName=body_font, fontSize=body_size, leading=body_leading, textColor=colors.HexColor("#1F2937"), spaceAfter=7, alignment=TA_JUSTIFY if options.document_style == "Legal Professional" else TA_LEFT),
		"h1": ParagraphStyle("LexH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=ink, spaceBefore=15, spaceAfter=8),
		"h2": ParagraphStyle("LexH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13.5, leading=17, textColor=accent, spaceBefore=12, spaceAfter=6),
		"h3": ParagraphStyle("LexH3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=ink, spaceBefore=9, spaceAfter=4),
		"bullet": ParagraphStyle("LexBullet", parent=styles["BodyText"], fontName=body_font, fontSize=body_size, leading=body_leading, leftIndent=18, firstLineIndent=-9, bulletIndent=7, spaceAfter=4),
		"code": ParagraphStyle("LexCode", parent=styles["Code"], fontName="Courier", fontSize=8.5, leading=11, backColor=colors.HexColor("#F1F5F9"), borderPadding=7, spaceBefore=5, spaceAfter=7),
		"table_cell": ParagraphStyle("LexTableCell", parent=styles["BodyText"], fontName=body_font, fontSize=9.4, leading=12.2, textColor=colors.HexColor("#1F2937"), spaceAfter=0, alignment=TA_LEFT),
		"table_header": ParagraphStyle("LexTableHeader", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=9.2, leading=11.5, textColor=ink, spaceAfter=0, alignment=TA_LEFT),
	}
	story = []
	logo_path = _brand_logo_path()
	if options.include_cover_page:
		if logo_path and os.path.exists(logo_path):
			story.append(Image(logo_path, width=0.58 * inch, height=0.58 * inch))
		story.extend([
			Spacer(1, 0.35 * inch),
			Paragraph("LEXOCRATES LEGAL SERVICES PVT. LTD.", ParagraphStyle("Kicker", fontName="Helvetica-Bold", fontSize=9, leading=12, textColor=accent, spaceAfter=16)),
			Paragraph(_pdf_inline(options.document_title), ParagraphStyle("CoverTitle", fontName="Helvetica-Bold", fontSize=25, leading=30, textColor=ink, spaceAfter=12)),
			Paragraph(_pdf_inline(metadata.get("subtitle") or "AI-assisted legal operations deliverable"), ParagraphStyle("CoverSubtitle", fontName="Helvetica", fontSize=12, leading=17, textColor=muted, spaceAfter=24)),
		])
		if options.include_metadata:
			story.append(_pdf_metadata_table(metadata, colors, Table, TableStyle, Paragraph, style_map["body"], inch))
		story.append(Spacer(1, 0.2 * inch))
		if options.confidentiality_label != "None":
			story.append(Paragraph(options.confidentiality_label.upper(), ParagraphStyle("Confidential", fontName="Helvetica-Bold", fontSize=9, textColor=colors.HexColor("#8A3B2F"), spaceBefore=12)))
		story.append(PageBreak())
	else:
		story.append(Paragraph(_pdf_inline(options.document_title), style_map["h1"]))
		if options.include_metadata:
			story.append(_pdf_metadata_table(metadata, colors, Table, TableStyle, Paragraph, style_map["body"], inch))
		story.append(Spacer(1, 8))

	for block in parse_markdown_blocks(markdown_text):
		kind = block["type"]
		if kind in {"h1", "h2", "h3"}:
			story.append(Paragraph(_pdf_inline(block["text"]), style_map[kind]))
		elif kind == "paragraph":
			story.append(Paragraph(_pdf_inline(block["text"]), style_map["body"]))
		elif kind in {"bullet", "number"}:
			marker = "•" if kind == "bullet" else f"{block['number']}."
			story.append(Paragraph(_pdf_inline(block["text"]), style_map["bullet"], bulletText=marker))
		elif kind == "code":
			story.append(Paragraph(html.escape(block["text"]).replace("\n", "<br/>"), style_map["code"]))
		elif kind == "rule":
			story.append(Spacer(1, 7))
		elif kind == "table":
			rows = [
				[
					Paragraph(_pdf_inline(cell), style_map["table_header"] if row_index == 0 else style_map["table_cell"])
					for cell in row
				]
				for row_index, row in enumerate(block["rows"])
			]
			if rows:
				columns = len(rows[0])
				proportions = [0.25, 0.2, 0.35, 0.2] if columns == 4 else [1 / columns] * columns
				table = Table(rows, colWidths=[doc.width * value for value in proportions], repeatRows=1, hAlign="LEFT")
				table.setStyle(TableStyle([
					("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
					("TEXTCOLOR", (0, 0), (-1, 0), ink),
					("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
					("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
					("VALIGN", (0, 0), (-1, -1), "TOP"),
					("LEFTPADDING", (0, 0), (-1, -1), 6),
					("RIGHTPADDING", (0, 0), (-1, -1), 6),
					("TOPPADDING", (0, 0), (-1, -1), 6),
					("BOTTOMPADDING", (0, 0), (-1, -1), 6),
				]))
				story.extend([table, Spacer(1, 8)])

	def decorate_page(canvas, document):
		width, height = page
		canvas.saveState()
		if document.page > (1 if options.include_cover_page else 0):
			if logo_path and os.path.exists(logo_path):
				canvas.drawImage(logo_path, 0.78 * inch, height - 0.58 * inch, width=0.24 * inch, height=0.24 * inch, preserveAspectRatio=True, mask="auto")
			canvas.setFont("Helvetica-Bold", 8.5)
			canvas.setFillColor(ink)
			canvas.drawString(1.08 * inch, height - 0.49 * inch, "LEXOCRATES LEGAL SERVICES")
			canvas.setStrokeColor(colors.HexColor("#D7DEE8"))
			canvas.line(0.78 * inch, height - 0.63 * inch, width - 0.78 * inch, height - 0.63 * inch)
		if options.confidentiality_label != "None":
			canvas.setFont("Helvetica-Bold", 7.5)
			canvas.setFillColor(colors.HexColor("#7A8798"))
			canvas.drawString(0.78 * inch, 0.42 * inch, options.confidentiality_label.upper())
			if not options.include_cover_page or document.page > 1:
				canvas.saveState()
				canvas.setFont("Helvetica-Bold", 34)
				canvas.setFillColor(colors.HexColor("#6B7280"))
				try:
					canvas.setFillAlpha(0.028)
				except Exception:
					canvas.setFillColor(colors.HexColor("#EEF1F5"))
				canvas.translate(width / 2, height / 2)
				canvas.rotate(42)
				canvas.drawCentredString(0, 0, options.confidentiality_label.upper())
				canvas.restoreState()
		if options.include_page_numbers:
			canvas.setFont("Helvetica", 8)
			canvas.setFillColor(colors.HexColor("#64748B"))
			canvas.drawRightString(width - 0.78 * inch, 0.42 * inch, f"Page {document.page}")
		canvas.restoreState()

	doc.build(story, onFirstPage=decorate_page, onLaterPages=decorate_page)
	value = buffer.getvalue()
	if not value.startswith(b"%PDF-"):
		raise frappe.ValidationError(_("Generated PDF failed signature validation."))
	return value


def generate_docx_bytes(markdown_text: str, metadata: dict, options: ExportOptions) -> bytes:
	try:
		from docx import Document
		from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
		from docx.enum.text import WD_ALIGN_PARAGRAPH
		from docx.oxml import OxmlElement
		from docx.oxml.ns import qn
		from docx.shared import Cm, Inches, Pt, RGBColor
	except ImportError:
		frappe.throw(_("DOCX export dependency is unavailable. Install python-docx."), frappe.ValidationError)

	doc = Document()
	section = doc.sections[0]
	if options.page_size == "A4":
		section.page_width, section.page_height = Cm(21), Cm(29.7)
	else:
		section.page_width, section.page_height = Inches(8.5), Inches(11)
	section.top_margin = Inches(0.82)
	section.bottom_margin = Inches(0.72)
	section.left_margin = Inches(0.82)
	section.right_margin = Inches(0.82)
	section.header_distance = Inches(0.32)
	section.footer_distance = Inches(0.32)

	ink = RGBColor(0x22, 0x34, 0x4F)
	accent = RGBColor(0x19, 0x76, 0xD2)
	muted = RGBColor(0x60, 0x70, 0x86)
	body_font = "Times New Roman" if options.document_style == "Legal Professional" else "Calibri"
	_configure_docx_styles(doc, body_font, ink, accent, Pt, qn)
	_configure_docx_header_footer(doc, section, metadata, options, ink, muted, Inches, Pt, qn, OxmlElement, WD_ALIGN_PARAGRAPH)

	if options.include_cover_page:
		logo_path = _brand_logo_path()
		p = doc.add_paragraph()
		p.paragraph_format.space_after = Pt(18)
		if logo_path and os.path.exists(logo_path):
			p.add_run().add_picture(logo_path, width=Inches(0.62))
		p = doc.add_paragraph()
		p.paragraph_format.space_after = Pt(10)
		run = p.add_run("LEXOCRATES LEGAL SERVICES PVT. LTD.")
		_set_docx_run(run, "Calibri", 9, accent, True, qn)
		p = doc.add_paragraph()
		p.paragraph_format.space_after = Pt(8)
		run = p.add_run(options.document_title)
		_set_docx_run(run, "Calibri", 24, ink, True, qn)
		p = doc.add_paragraph()
		p.paragraph_format.space_after = Pt(18)
		run = p.add_run(metadata.get("subtitle") or "AI-assisted legal operations deliverable")
		_set_docx_run(run, "Calibri", 12, muted, False, qn)
		if options.include_metadata:
			_add_docx_metadata_table(doc, metadata, Inches, Pt, qn, OxmlElement, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT)
		if options.confidentiality_label != "None":
			p = doc.add_paragraph()
			p.paragraph_format.space_before = Pt(14)
			run = p.add_run(options.confidentiality_label.upper())
			_set_docx_run(run, "Calibri", 9, RGBColor(0x8A, 0x3B, 0x2F), True, qn)
		doc.add_page_break()
	else:
		p = doc.add_paragraph(style="Heading 1")
		p.add_run(options.document_title)
		if options.include_metadata:
			_add_docx_metadata_table(doc, metadata, Inches, Pt, qn, OxmlElement, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT)

	for block in parse_markdown_blocks(markdown_text):
		kind = block["type"]
		if kind in {"h1", "h2", "h3"}:
			p = doc.add_paragraph(style={"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3"}[kind])
			_add_docx_inline(p, block["text"], body_font, qn)
		elif kind == "paragraph":
			p = doc.add_paragraph()
			_add_docx_inline(p, block["text"], body_font, qn)
		elif kind in {"bullet", "number"}:
			p = doc.add_paragraph(style="List Bullet" if kind == "bullet" else "List Number")
			p.paragraph_format.left_indent = Inches(0.5)
			p.paragraph_format.first_line_indent = Inches(-0.25)
			p.paragraph_format.space_after = Pt(4)
			_add_docx_inline(p, block["text"], body_font, qn)
		elif kind == "code":
			p = doc.add_paragraph()
			p.paragraph_format.left_indent = Inches(0.25)
			p.paragraph_format.space_after = Pt(7)
			run = p.add_run(block["text"])
			_set_docx_run(run, "Courier New", 9, RGBColor(0x33, 0x41, 0x55), False, qn)
		elif kind == "rule":
			p = doc.add_paragraph()
			p.paragraph_format.space_after = Pt(6)
			pPr = p._p.get_or_add_pPr()
			pBdr = OxmlElement("w:pBdr")
			bottom = OxmlElement("w:bottom")
			bottom.set(qn("w:val"), "single")
			bottom.set(qn("w:sz"), "4")
			bottom.set(qn("w:color"), "CBD5E1")
			pBdr.append(bottom)
			pPr.append(pBdr)
		elif kind == "table":
			_add_docx_table(doc, block["rows"], body_font, Inches, Pt, qn, OxmlElement, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT)

	doc.core_properties.title = options.document_title
	doc.core_properties.subject = f"LPO Job {metadata.get('job_id') or ''}"
	doc.core_properties.author = "Lexocrates Legal Services Pvt. Ltd."
	doc.core_properties.last_modified_by = "Lexocrates Document Export Engine"
	doc.core_properties.comments = "Generated natively inside the Lexocrates Frappe application."
	buffer = io.BytesIO()
	doc.save(buffer)
	value = buffer.getvalue()
	if not value.startswith(b"PK") or b"word/document.xml" not in value:
		raise frappe.ValidationError(_("Generated DOCX failed package validation."))
	return value


def parse_markdown_blocks(text: str) -> list[dict]:
	lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
	blocks = []
	paragraph = []
	i = 0

	def flush_paragraph():
		if paragraph:
			blocks.append({"type": "paragraph", "text": " ".join(item.strip() for item in paragraph).strip()})
			paragraph.clear()

	while i < len(lines):
		line = lines[i]
		stripped = line.strip()
		if stripped.startswith("```"):
			flush_paragraph()
			code_lines = []
			i += 1
			while i < len(lines) and not lines[i].strip().startswith("```"):
				code_lines.append(lines[i])
				i += 1
			blocks.append({"type": "code", "text": "\n".join(code_lines)})
		elif stripped.startswith("|") and i + 1 < len(lines) and _is_markdown_table_separator(lines[i + 1]):
			flush_paragraph()
			rows = [_split_markdown_table_row(line)]
			i += 2
			while i < len(lines) and lines[i].strip().startswith("|"):
				rows.append(_split_markdown_table_row(lines[i]))
				i += 1
			i -= 1
			max_cols = max(len(row) for row in rows)
			blocks.append({"type": "table", "rows": [row + [""] * (max_cols - len(row)) for row in rows]})
		elif not stripped:
			flush_paragraph()
		elif re.match(r"^#{1,3}\s+", stripped):
			flush_paragraph()
			level = len(stripped) - len(stripped.lstrip("#"))
			blocks.append({"type": f"h{level}", "text": stripped[level:].strip()})
		elif re.match(r"^[-*+]\s+", stripped):
			flush_paragraph()
			blocks.append({"type": "bullet", "text": re.sub(r"^[-*+]\s+", "", stripped)})
		elif re.match(r"^\d+[.)]\s+", stripped):
			flush_paragraph()
			match = re.match(r"^(\d+)[.)]\s+(.*)", stripped)
			blocks.append({"type": "number", "number": int(match.group(1)), "text": match.group(2)})
		elif re.match(r"^([-*_])\1\1+$", stripped.replace(" ", "")):
			flush_paragraph()
			blocks.append({"type": "rule"})
		else:
			paragraph.append(stripped)
		i += 1
	flush_paragraph()
	return blocks


def _pdf_metadata_table(metadata, colors, Table, TableStyle, Paragraph, body_style, inch):
	rows = []
	for label, key in (("Client", "client_name"), ("Matter", "matter_title"), ("Job", "job_label"), ("Version", "version_label"), ("Generated", "generated_label"), ("Prepared by", "generated_by")):
		value = metadata.get(key)
		if value:
			rows.append([Paragraph(f"<b>{html.escape(label)}</b>", body_style), Paragraph(_pdf_inline(str(value)), body_style)])
	table = Table(rows, colWidths=[1.15 * inch, 4.95 * inch], hAlign="LEFT")
	table.setStyle(TableStyle([
		("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")),
		("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D7DEE8")),
		("VALIGN", (0, 0), (-1, -1), "TOP"),
		("LEFTPADDING", (0, 0), (-1, -1), 7),
		("RIGHTPADDING", (0, 0), (-1, -1), 7),
		("TOPPADDING", (0, 0), (-1, -1), 6),
		("BOTTOMPADDING", (0, 0), (-1, -1), 6),
	]))
	return table


def _configure_docx_styles(doc, body_font, ink, accent, Pt, qn):
	normal = doc.styles["Normal"]
	normal.font.name = body_font
	normal._element.rPr.rFonts.set(qn("w:ascii"), body_font)
	normal._element.rPr.rFonts.set(qn("w:hAnsi"), body_font)
	normal.font.size = Pt(11)
	normal.paragraph_format.space_after = Pt(6)
	normal.paragraph_format.line_spacing = 1.10
	for name, size, color, before, after in (
		("Heading 1", 16, ink, 16, 8),
		("Heading 2", 13, accent, 12, 6),
		("Heading 3", 12, ink, 8, 4),
	):
		style = doc.styles[name]
		style.font.name = "Calibri"
		style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
		style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
		style.font.size = Pt(size)
		style.font.bold = True
		style.font.color.rgb = color
		style.paragraph_format.space_before = Pt(before)
		style.paragraph_format.space_after = Pt(after)
		style.paragraph_format.keep_with_next = True


def _configure_docx_header_footer(doc, section, metadata, options, ink, muted, Inches, Pt, qn, OxmlElement, align):
	header = section.header
	p = header.paragraphs[0]
	p.alignment = align.LEFT
	logo_path = _brand_logo_path()
	if logo_path and os.path.exists(logo_path):
		p.add_run().add_picture(logo_path, width=Inches(0.22))
	run = p.add_run("  LEXOCRATES LEGAL SERVICES")
	_set_docx_run(run, "Calibri", 8.5, ink, True, qn)
	if options.confidentiality_label != "None":
		run = p.add_run(f"   |   {options.confidentiality_label.upper()}")
		_set_docx_run(run, "Calibri", 8, muted, True, qn)
	footer = section.footer
	p = footer.paragraphs[0]
	p.alignment = align.RIGHT
	if options.include_page_numbers:
		run = p.add_run("Page ")
		_set_docx_run(run, "Calibri", 8, muted, False, qn)
		fld_char = OxmlElement("w:fldChar")
		fld_char.set(qn("w:fldCharType"), "begin")
		instr = OxmlElement("w:instrText")
		instr.set(qn("xml:space"), "preserve")
		instr.text = "PAGE"
		end = OxmlElement("w:fldChar")
		end.set(qn("w:fldCharType"), "end")
		run._r.extend([fld_char, instr, end])
	else:
		run = p.add_run(metadata.get("job_id") or "Lexocrates Deliverable")
		_set_docx_run(run, "Calibri", 8, muted, False, qn)


def _add_docx_metadata_table(doc, metadata, Inches, Pt, qn, OxmlElement, table_align, cell_align):
	items = [(label, metadata.get(key)) for label, key in (("Client", "client_name"), ("Matter", "matter_title"), ("Job", "job_label"), ("Version", "version_label"), ("Generated", "generated_label"), ("Prepared by", "generated_by")) if metadata.get(key)]
	table = doc.add_table(rows=len(items), cols=2)
	table.style = "Table Grid"
	table.alignment = table_align.LEFT
	table.autofit = False
	_set_docx_table_geometry(table, [1.3, 5.2], qn, OxmlElement)
	for index, (label, value) in enumerate(items):
		for cell in table.rows[index].cells:
			cell.vertical_alignment = cell_align.CENTER
		cell = table.rows[index].cells[0]
		_set_cell_fill(cell, "F2F4F7", qn, OxmlElement)
		p = cell.paragraphs[0]
		p.paragraph_format.space_after = Pt(2)
		run = p.add_run(label)
		_set_docx_run(run, "Calibri", 9.5, None, True, qn)
		p = table.rows[index].cells[1].paragraphs[0]
		p.paragraph_format.space_after = Pt(2)
		run = p.add_run(str(value))
		_set_docx_run(run, "Calibri", 9.5, None, False, qn)
	return table


def _add_docx_table(doc, rows, body_font, Inches, Pt, qn, OxmlElement, table_align, cell_align):
	if not rows:
		return
	columns = max(len(row) for row in rows)
	table = doc.add_table(rows=len(rows), cols=columns)
	table.style = "Table Grid"
	table.alignment = table_align.LEFT
	table.autofit = False
	_set_docx_table_geometry(table, [6.5 / columns] * columns, qn, OxmlElement)
	for row_index, values in enumerate(rows):
		for col_index in range(columns):
			cell = table.rows[row_index].cells[col_index]
			cell.vertical_alignment = cell_align.CENTER
			if row_index == 0:
				_set_cell_fill(cell, "E8EEF5", qn, OxmlElement)
			p = cell.paragraphs[0]
			p.paragraph_format.space_after = Pt(2)
			run = p.add_run(values[col_index] if col_index < len(values) else "")
			_set_docx_run(run, "Calibri" if row_index == 0 else body_font, 9.5, None, row_index == 0, qn)
	doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _set_docx_table_geometry(table, widths_in, qn, OxmlElement):
	total = round(sum(widths_in) * 1440)
	tbl_pr = table._tbl.tblPr
	tbl_w = tbl_pr.find(qn("w:tblW"))
	if tbl_w is None:
		tbl_w = OxmlElement("w:tblW")
		tbl_pr.append(tbl_w)
	tbl_w.set(qn("w:type"), "dxa")
	tbl_w.set(qn("w:w"), str(total))
	tbl_ind = OxmlElement("w:tblInd")
	tbl_ind.set(qn("w:type"), "dxa")
	tbl_ind.set(qn("w:w"), "120")
	tbl_pr.append(tbl_ind)
	for index, width in enumerate(widths_in):
		width_dxa = round(width * 1440)
		if index < len(table._tbl.tblGrid.gridCol_lst):
			table._tbl.tblGrid.gridCol_lst[index].set(qn("w:w"), str(width_dxa))
		for row in table.rows:
			cell = row.cells[index]
			tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
			if tc_w is None:
				tc_w = OxmlElement("w:tcW")
				cell._tc.get_or_add_tcPr().append(tc_w)
			tc_w.set(qn("w:type"), "dxa")
			tc_w.set(qn("w:w"), str(width_dxa))


def _set_cell_fill(cell, fill, qn, OxmlElement):
	shading = OxmlElement("w:shd")
	shading.set(qn("w:fill"), fill)
	cell._tc.get_or_add_tcPr().append(shading)


def _add_docx_inline(paragraph, text, font_name, qn):
	for value, bold, italic, code in _inline_tokens(text):
		run = paragraph.add_run(value)
		_set_docx_run(run, "Courier New" if code else font_name, 10 if code else None, None, bold, qn)
		run.italic = italic


def _set_docx_run(run, font_name, size, color, bold, qn):
	run.font.name = font_name
	run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font_name)
	run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font_name)
	if size is not None:
		from docx.shared import Pt

		run.font.size = Pt(size)
	if color is not None:
		run.font.color.rgb = color
	if bold is not None:
		run.bold = bold


def _inline_tokens(text: str):
	pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|(?<!\*)\*[^*]+?\*)")
	position = 0
	for match in pattern.finditer(text or ""):
		if match.start() > position:
			yield text[position:match.start()], False, False, False
		token = match.group(0)
		if token.startswith("**"):
			yield token[2:-2], True, False, False
		elif token.startswith("`"):
			yield token[1:-1], False, False, True
		else:
			yield token[1:-1], False, True, False
		position = match.end()
	if position < len(text or ""):
		yield text[position:], False, False, False


def _pdf_inline(text: str) -> str:
	parts = []
	for value, bold, italic, code in _inline_tokens(str(text or "")):
		value = html.escape(value).replace("\n", "<br/>")
		if code:
			parts.append(f"<font name='Courier'>{value}</font>")
		elif bold:
			parts.append(f"<b>{value}</b>")
		elif italic:
			parts.append(f"<i>{value}</i>")
		else:
			parts.append(value)
	return "".join(parts)


def _is_markdown_table_separator(line: str) -> bool:
	cells = _split_markdown_table_row(line)
	return bool(cells) and all(re.match(r"^:?-{3,}:?$", cell.replace(" ", "")) for cell in cells)


def _split_markdown_table_row(line: str) -> list[str]:
	return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _brand_logo_path() -> str:
	return frappe.get_app_path("lex", "public", "images", "lexocrates-mark-dark.png")


def safe_export_filename(job_id: str, title: str, version: int, extension: str) -> str:
	stem = re.sub(r"[^A-Za-z0-9_-]+", "_", title or "Deliverable").strip("_")[:70] or "Deliverable"
	return f"Delivery_{job_id}_{stem}_v{version}.{extension.lower()}"


def export_metadata(job, matter, *, title: str, version: int, generated_by: str) -> dict:
	client_name = frappe.db.get_value("Customer", job.customer, "customer_name") or job.customer
	return {
		"title": title,
		"subtitle": f"{job.job_type or 'Legal Operations'} - Final Deliverable",
		"client_name": client_name,
		"matter_title": matter.matter_title or matter.name,
		"job_id": job.name,
		"job_label": f"{job.name} - {job.job_title}",
		"version_label": f"v{version}.0",
		"generated_by": generated_by,
		"generated_label": datetime.now().astimezone().strftime("%d %B %Y, %H:%M %Z"),
	}
